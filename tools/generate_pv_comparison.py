#!/usr/bin/env python3
"""Generate a Word document comparing PropertyValue implementations across building blocks."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

def set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(shd)

def main():
    doc = Document()

    # Title
    title = doc.add_heading('PropertyValue Implementation Comparison', level=1)
    doc.add_paragraph(
        'Comparison of schema:PropertyValue properties across CDIF building blocks '
        'where PropertyValue is used as the type for a schema property. '
        'Generated from schema.yaml source files.'
    )

    # Define the contexts (columns)
    contexts = [
        'identifier',
        'variableMeasured',
        'additionalProperty',
        'cdifInstanceVariable\n(extends variableMeasured)',
        'bioschemas\nParameterValue',
    ]

    # Define all properties found across all contexts
    # Each row: (property_name, description, then per-context: (cardinality, notes))
    # Cardinality: "1" = required single, "0..1" = optional single, "1..*" = required array,
    #              "0..*" = optional array, "-" = not present

    rows = [
        ('@type',
         'Type declaration (must contain schema:PropertyValue)',
         [
             ('1..*\nrequired', 'contains: PropertyValue'),
             ('1..*\nrequired', 'contains: PropertyValue'),
             ('1..*\nrequired', 'contains: PropertyValue'),
             ('1..*\nrequired', 'contains: cdi:InstanceVariable\n(+ PropertyValue via variableMeasured)'),
             ('1..*\nrequired', 'contains: PropertyValue'),
         ]),
        ('@id',
         'URI identifier for this node',
         [
             ('-', ''),
             ('0..1', 'string'),
             ('-', ''),
             ('-', '(inherited from variableMeasured)'),
             ('0..1', 'string'),
         ]),
        ('schema:name',
         'Human-readable label',
         [
             ('-', ''),
             ('1\nrequired', 'string'),
             ('1\nrequired', 'string'),
             ('-', '(inherited: required via variableMeasured)'),
             ('1\nrequired', 'string'),
         ]),
        ('schema:description',
         'Textual description',
         [
             ('-', ''),
             ('0..1', 'string, default: "missing"'),
             ('-', ''),
             ('-', '(inherited from variableMeasured)'),
             ('-', ''),
         ]),
        ('schema:alternateName',
         'Alternative names',
         [
             ('-', ''),
             ('0..*', 'array of strings'),
             ('-', ''),
             ('-', '(inherited from variableMeasured)'),
             ('-', ''),
         ]),
        ('schema:propertyID',
         'Identifier for the property concept',
         [
             ('0..1', 'string\n(identifier scheme name)'),
             ('0..*', 'array of: string | {@id} | DefinedTerm'),
             ('1..*\nrequired', 'array of: string | {@id} | DefinedTerm;\nminItems: 1'),
             ('-', '(inherited from variableMeasured)'),
             ('0..1', 'string | {@id}\n(link to FormalParameter)'),
         ]),
        ('schema:value',
         'The property value',
         [
             ('0..1\n(conditional)', 'string;\nrequired if no schema:url'),
             ('-', ''),
             ('1\nrequired', 'string | number | boolean | object'),
             ('-', ''),
             ('0..1', 'string | number | boolean'),
         ]),
        ('schema:url',
         'Web-resolvable URL',
         [
             ('0..1\n(conditional)', 'string (uri format);\nrequired if no schema:value'),
             ('0..1', 'string (uri) | LabeledLink'),
             ('-', ''),
             ('-', '(inherited from variableMeasured)'),
             ('-', ''),
         ]),
        ('schema:unitText',
         'Unit of measurement as text',
         [
             ('-', ''),
             ('0..1', 'string'),
             ('0..1', 'string'),
             ('-', '(inherited from variableMeasured)'),
             ('0..1', 'string'),
         ]),
        ('schema:unitCode',
         'URI or code for unit of measure',
         [
             ('-', ''),
             ('0..1', 'string | {@id} | DefinedTerm'),
             ('0..1', 'string | DefinedTerm'),
             ('-', '(inherited from variableMeasured)'),
             ('-', ''),
         ]),
        ('schema:measurementTechnique',
         'How values were obtained',
         [
             ('-', ''),
             ('0..1', 'string | {@id} | DefinedTerm'),
             ('-', ''),
             ('-', '(inherited from variableMeasured)'),
             ('-', ''),
         ]),
        ('schema:minValue',
         'Minimum numeric value',
         [
             ('-', ''),
             ('0..1', 'number'),
             ('-', ''),
             ('-', '(inherited from variableMeasured)'),
             ('-', ''),
         ]),
        ('schema:maxValue',
         'Maximum numeric value',
         [
             ('-', ''),
             ('0..1', 'number'),
             ('-', ''),
             ('-', '(inherited from variableMeasured)'),
             ('-', ''),
         ]),
        ('cdi:identifier',
         'DDI-CDI identifier',
         [
             ('-', ''),
             ('-', ''),
             ('-', ''),
             ('0..1', 'string'),
             ('-', ''),
         ]),
        ('cdi:physicalDataType',
         'Physical data type',
         [
             ('-', ''),
             ('-', ''),
             ('-', ''),
             ('0..*', 'array of: string | {@id} | DefinedTerm'),
             ('-', ''),
         ]),
        ('cdi:intendedDataType',
         'Intended data type (XML Schema)',
         [
             ('-', ''),
             ('-', ''),
             ('-', ''),
             ('0..1', 'string'),
             ('-', ''),
         ]),
        ('cdi:role',
         'Variable role in data structure',
         [
             ('-', ''),
             ('-', ''),
             ('-', ''),
             ('0..1', 'enum: Measure, Attribute,\nDimension, Descriptor,\nReferenceValue'),
             ('-', ''),
         ]),
        ('cdi:describedUnitOfMeasure',
         'Structured unit from vocabulary',
         [
             ('-', ''),
             ('-', ''),
             ('-', ''),
             ('0..1', 'DefinedTerm'),
             ('-', ''),
         ]),
        ('cdi:simpleUnitOfMeasure',
         'Simple unit reference',
         [
             ('-', ''),
             ('-', ''),
             ('-', ''),
             ('0..1', 'string | {@id} | DefinedTerm'),
             ('-', ''),
         ]),
        ('cdi:uses',
         'Concept references',
         [
             ('-', ''),
             ('-', ''),
             ('-', ''),
             ('0..*', 'array of: string | {@id} | DefinedTerm'),
             ('-', ''),
         ]),
        ('cdi:name',
         'DDI-CDI Concept name',
         [
             ('-', ''),
             ('-', ''),
             ('-', ''),
             ('0..1', 'string'),
             ('-', ''),
         ]),
        ('cdi:displayLabel',
         'Human-readable display label',
         [
             ('-', ''),
             ('-', ''),
             ('-', ''),
             ('0..1', 'string'),
             ('-', ''),
         ]),
        ('cdi:qualifies',
         'Reference to another instance variable',
         [
             ('-', ''),
             ('-', ''),
             ('-', ''),
             ('0..1', 'object with @id'),
             ('-', ''),
         ]),
    ]

    # Create table: property + description + 5 context columns = 7 columns
    num_cols = 2 + len(contexts)
    table = doc.add_table(rows=1, cols=num_cols, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    hdr = table.rows[0]
    headers = ['Property', 'Description'] + contexts
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(8)
        set_cell_shading(cell, '4472C4')
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for prop_name, description, context_data in rows:
        row = table.add_row()
        # Property name
        cell0 = row.cells[0]
        run0 = cell0.paragraphs[0].add_run(prop_name)
        run0.bold = True
        run0.font.size = Pt(8)

        # Description
        cell1 = row.cells[1]
        run1 = cell1.paragraphs[0].add_run(description)
        run1.font.size = Pt(7)

        # Context columns
        for j, (cardinality, notes) in enumerate(context_data):
            cell = row.cells[2 + j]
            p = cell.paragraphs[0]

            if cardinality == '-':
                run = p.add_run('-')
                run.font.size = Pt(7)
                run.font.color.rgb = RGBColor(180, 180, 180)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_cell_shading(cell, 'F2F2F2')
            else:
                # Cardinality
                run_c = p.add_run(cardinality)
                run_c.font.size = Pt(7)
                if 'required' in cardinality:
                    run_c.bold = True

                # Notes on new line
                if notes:
                    run_n = p.add_run('\n' + notes)
                    run_n.font.size = Pt(6.5)
                    run_n.font.color.rgb = RGBColor(100, 100, 100)

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(1.3)
        row.cells[1].width = Inches(1.5)
        for j in range(len(contexts)):
            row.cells[2 + j].width = Inches(1.5)

    # Add notes section
    doc.add_paragraph()
    doc.add_heading('Notes', level=2)

    notes_items = [
        'identifier: Used for schema:identifier property. PropertyValue encodes structured identifiers (DOI, ARK, ORCID, etc.). Requires either schema:value or schema:url.',
        'variableMeasured: Used for schema:variableMeasured property. PropertyValue describes a measured variable in a dataset with its name, units, range, and measurement technique.',
        'additionalProperty: Used for schema:additionalProperty property. PropertyValue defines soft-typed key-value properties. Requires schema:name and schema:value; schema:propertyID is required (minItems: 1) for semantic identification.',
        'cdifInstanceVariable: Extends variableMeasured via allOf composition with DDI-CDI InstanceVariable properties. Inherits all variableMeasured properties and adds CDI-specific fields for data structure roles, types, and cross-variable references.',
        'bioschemas ParameterValue: Used for bios:parameterValue in lab process descriptions. Records actual instrument settings or measurement conditions (temperature, flow rate, etc.) during a process execution.',
        'Cardinality notation: "1" = exactly one (required), "0..1" = optional single value, "1..*" = required array (one or more), "0..*" = optional array (zero or more), "-" = not present in this context.',
    ]

    for item in notes_items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(9)

    # Building block source files
    doc.add_heading('Source Files', level=2)
    sources = [
        ('identifier', '_sources/schemaorgProperties/identifier/schema.yaml'),
        ('variableMeasured', '_sources/schemaorgProperties/variableMeasured/schema.yaml'),
        ('additionalProperty', '_sources/schemaorgProperties/additionalProperty/schema.yaml'),
        ('cdifInstanceVariable', '_sources/cdifProperties/cdifInstanceVariable/schema.yaml'),
        ('bioschemas ParameterValue', '_sources/bioschemasProperties/cdifBioschemasProperties/schema.yaml (ParameterValue $def)'),
    ]
    for name, path in sources:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{name}: ')
        run.bold = True
        run.font.size = Pt(9)
        run2 = p.add_run(path)
        run2.font.size = Pt(9)

    # Save
    output_dir = os.path.dirname(__file__)
    output_path = os.path.join(output_dir, '..', 'PropertyValue_Comparison.docx')
    output_path = os.path.normpath(output_path)
    doc.save(output_path)
    print(f'Saved to: {output_path}')


if __name__ == '__main__':
    main()
